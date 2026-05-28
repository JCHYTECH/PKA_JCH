from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("JCH_Inbox/03_PROJECTS/08_VETALYX/clinical")
DOCX = OUT / "2026-05-26_VETALYX_Rapport_final_tests_rapides_allergies_chien_chat.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(88, 96, 105)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
BORDER = "C8D1DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def paragraph(doc, text="", style=None, size=11, bold=False, color=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    for run in p.runs:
        set_run(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="B7C7D8", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, size=11.5, bold=True, color=INK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=INK)
    paragraph(doc, "", after=4)


def add_table(doc, columns, rows, widths):
    table = doc.add_table(rows=1, cols=len(columns))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        set_cell_shading(hdr[i], LIGHT_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(col)
        set_run(r, size=9.4, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(val)
            set_run(r, size=9.1, color=RGBColor(20, 20, 20))
    paragraph(doc, "", after=6)
    return table


def add_source_note(doc, source_ids):
    p = paragraph(doc, "Sources principales: " + ", ".join(source_ids), size=8.6, color=MUTED, after=7, before=0)
    p.paragraph_format.line_spacing = 1.0


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "VETALYX | Rapid allergy tests - clinical product specification"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Confidentiel - version de travail partenaire - 26 mai 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run(run, size=8.5, color=MUTED)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def masthead(doc):
    paragraph(doc, "VETALYX", size=10.5, bold=True, color=MUTED, after=2)
    title = paragraph(doc, "Rapport final - specification des tests rapides allergie chien/chat", size=23, bold=True, color=RGBColor(0, 0, 0), after=4)
    title.paragraph_format.line_spacing = 1.0
    paragraph(doc, "Panels 6 lignes pour IgE respiratoire, dermique et alimentaire - version partenaire technique/scientifique", size=13, color=MUTED, after=14)
    meta = [
        ("Date", "26 mai 2026"),
        ("Auteur", "Dobby - orchestration PKA_JCH ; revue clinique Vasco ; revue scientifique Clio"),
        ("Objet", "Definition des lignes d'allergenes pour tests rapides chien et chat"),
        ("Statut", "Rapport distribuable v1 - a valider avec KOL dermatologie veterinaire et fournisseur"),
    ]
    for label, value in meta:
        p = paragraph(doc, after=2)
        r1 = p.add_run(f"{label}: ")
        set_run(r1, bold=True)
        r2 = p.add_run(value)
        set_run(r2)
    paragraph(doc, "", after=8)


def make_doc():
    doc = Document()
    configure_doc(doc)
    masthead(doc)

    add_callout(
        doc,
        "Decision centrale",
        "Avec seulement 6 lignes, les tests rapides doivent utiliser des lignes mixtes par familles allergeniques. "
        "Une ligne positive signifie une sensibilisation IgE au groupe, pas l'identification de l'espece allergenique exacte.",
    )
    paragraph(
        doc,
        "La priorite P0 recommandee est un test environnemental/dermatologique chien et chat. Les panels alimentaires doivent rester en P2, "
        "avec un claim limite a l'aide a l'orientation d'un regime d'eviction.",
    )

    heading(doc, "1. Regles d'interpretation", 1)
    bullets(
        doc,
        [
            "Les lignes mixtes sont recommandees pour les acariens, pollens et moisissures afin d'augmenter le rendement clinique dans un format 6 lignes.",
            "Un resultat positif doit etre interprete dans le contexte clinique: prurit, saisonnalite, dermatite, controle parasitaire, infections secondaires et historique alimentaire.",
            "Pour l'immunotherapie specifique, un panel laboratoire plus detaille peut etre necessaire apres le test rapide.",
            "Les IgE alimentaires seules ne diagnostiquent pas une allergie alimentaire: le regime d'eviction/provocation reste la reference clinique.",
        ],
    )

    heading(doc, "2. Panels chien", 1)
    heading(doc, "2.1 Chien - IgE respiratoire / aeroallergenes", 2)
    add_source_note(doc, ["S1", "S2", "S3", "S4", "S5"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "House dust mites", "Dermatophagoides farinae + D. pteronyssinus", "Groupe le plus robuste en dermatite atopique canine ; D. farinae seul est incomplet."),
            ("2", "Storage mites", "Tyrophagus putrescentiae + Acarus siro + Lepidoglyphus destructor", "Frequent en environnement domestique, aliments secs et zones humides ; utile en screening EU."),
            ("3", "Grass pollen mix", "Phleum/Timothy + Lolium + Dactylis + Festuca + Poa", "Graminees tres pertinentes en Europe ; saisonnalite exploitable cliniquement."),
            ("4", "Weed pollen mix", "Artemisia + Ambrosia + Plantago + Parietaria/Urticaceae", "Mauvaises herbes importantes ; Parietaria prioritaire Sud/Mediterranee."),
            ("5", "Tree pollen mix", "Betula + Alnus + Corylus + Quercus + Fraxinus/Platanus", "Arbres europeens a forte variabilite regionale ; option olive/cypress au Sud."),
            ("6", "Mould mix", "Alternaria + Cladosporium + Aspergillus + Penicillium", "Pertinence variable mais utile en ligne groupee."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "2.2 Chien - IgE dermique / dermatite allergique", 2)
    add_source_note(doc, ["S1", "S2", "S3", "S5"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "Flea allergy", "Ctenocephalides felis / flea saliva", "Cause majeure de dermatite allergique ; ligne prioritaire meme si l'interpretation IgE n'est pas parfaite."),
            ("2", "House dust mites", "D. farinae + D. pteronyssinus", "Fort rendement pour dermatite atopique canine."),
            ("3", "Storage mites", "Tyrophagus + Acarus + Lepidoglyphus", "Complete la ligne HDM ; utile en Europe."),
            ("4", "Malassezia", "Malassezia pachydermatis", "Sensibilisation possible chez chiens atopiques avec dermatite/otite ; a interpreter avec infection active."),
            ("5", "Atopic pollen mix", "Grass mix + Artemisia/Plantago/Parietaria selon region", "Couvre les pollens cutaneo-atopiques les plus actionnables."),
            ("6", "Mould mix", "Alternaria + Cladosporium + Aspergillus + Penicillium", "Complement environnemental en dermatite chronique ou saisonniere."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "2.3 Chien - IgE alimentaire", 2)
    add_source_note(doc, ["S6", "S7", "S8"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "Beef", "Proteines de boeuf / BSA si disponible", "Source alimentaire frequemment rapportee chez le chien."),
            ("2", "Dairy", "Proteines de lait de vache", "Allergene alimentaire frequent ; exposition courante."),
            ("3", "Chicken", "Proteines de poulet", "Ingredient tres expose dans les regimes canins."),
            ("4", "Wheat/cereal", "Ble +/- cereales selon contrainte technique", "Source rapportee ; utile pour orienter l'eviction."),
            ("5", "Egg", "Proteines de blanc d'oeuf", "Pertinent dans les listes de sources alimentaires courantes."),
            ("6", "Soy", "Proteines de soja", "Pertinence moyenne mais utile si ligne disponible."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "3. Panels chat", 1)
    heading(doc, "3.1 Chat - IgE respiratoire / aeroallergenes", 2)
    add_source_note(doc, ["S5", "S9", "S10"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "House dust mites", "D. farinae + D. pteronyssinus", "HDM ressortent souvent dans les reactions IgE/IDT feline ; interpretation plus prudente que chien."),
            ("2", "Storage mites", "Tyrophagus + Acarus + Lepidoglyphus", "Complement pertinent en environnement interieur et alimentation seche."),
            ("3", "Grass pollen mix", "Timothy/Phleum + Lolium + Dactylis + Poa", "Couverture pollens de base en Europe."),
            ("4", "Regional tree/weed mix", "Betula/Alnus/Corylus ou olive/cypress/Parietaria selon zone", "A adapter geographiquement ; valeur surtout contextuelle."),
            ("5", "Mould mix", "Alternaria + Cladosporium + Aspergillus + Penicillium", "Signal variable mais utile en screening respiratoire/environnemental."),
            ("6", "Indoor insects", "Cockroach +/- moustique/insect mix si disponible", "Interieur urbain ; pour asthme, doit rester un support et non un diagnostic."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "3.2 Chat - IgE dermique / dermatite allergique", 2)
    add_source_note(doc, ["S9", "S10", "S5"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "Flea allergy", "Ctenocephalides felis / flea saliva", "Priorite clinique feline ; toujours a exclure/traiter."),
            ("2", "House dust mites", "D. farinae + D. pteronyssinus", "Pertinent dans FASS/non-flea non-food hypersensitivity."),
            ("3", "Storage mites", "Tyrophagus + Acarus + Lepidoglyphus", "Ajoute une couverture interieur/alimentation seche."),
            ("4", "Regional pollen mix", "Graminees + weed/tree mix regional", "Aide si saisonnalite ou environnement compatible."),
            ("5", "Mould mix", "Alternaria + Cladosporium + Aspergillus + Penicillium", "Complement environnemental prudent."),
            ("6", "Biting insects", "Moustique/insect bite mix ; sinon cockroach", "Utile si disponible ; sinon ligne insectes indoor."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "3.3 Chat - IgE alimentaire", 2)
    add_source_note(doc, ["S6", "S7", "S8"])
    add_table(
        doc,
        ["Ligne", "Nom affiche", "Composition recommandee", "Rationale"],
        [
            ("1", "Beef", "Proteines de boeuf", "Source alimentaire frequemment rapportee."),
            ("2", "Fish", "Fish mix : poisson blanc + saumon/thon selon exposition", "Important chez le chat du fait des regimes courants."),
            ("3", "Chicken", "Proteines de poulet", "Exposition alimentaire tres frequente."),
            ("4", "Dairy", "Proteines de lait de vache", "Source rapportee ; ligne interpretable seulement en support."),
            ("5", "Egg", "Proteines d'oeuf", "Pertinent mais non diagnostique seul."),
            ("6", "Wheat / soy", "Ble + soja groupes, ou ligne separee selon contrainte", "Compromis si seulement 6 lignes."),
        ],
        [650, 1700, 3300, 3710],
    )

    heading(doc, "4. Claims recommandes", 1)
    add_callout(
        doc,
        "Claim autorise recommande",
        "Test rapide de sensibilisation IgE a des groupes d'allergenes frequents, destine a appuyer l'orientation clinique par le veterinaire.",
    )
    paragraph(doc, "Claims a eviter : diagnostic alimentaire rapide ; detecte les intolerances ; identifie la cause de l'allergie ; remplace le regime d'eviction ; remplace le diagnostic veterinaire.")

    heading(doc, "5. Decision fournisseur", 1)
    bullets(
        doc,
        [
            "Demander a DIXUN de confirmer la faisabilite technique de lignes mixtes standardisees par famille allergenique.",
            "Exiger une validation chien/chat separee, avec matrice positive/negative et reproductibilite par ligne.",
            "Clarifier si les extraits sont natifs, recombinants ou mixtes, et obtenir les controles qualite lot a lot.",
            "Prevoir une option regionale Nord/Ouest Europe et Mediterranee pour les pollens.",
        ],
    )

    heading(doc, "6. Bibliographie et sources", 1)
    sources = [
        ("S1", "ICADA canine atopic dermatitis treatment guidelines", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4537558/"),
        ("S2", "Canine atopic dermatitis diagnosis and allergen identification guidelines", "https://bmcvetres.biomedcentral.com/articles/10.1186/s12917-015-0515-5"),
        ("S3", "Efficacy of diagnostic testing for allergen sensitization in canine atopic dermatitis - systematic review", "https://pmc.ncbi.nlm.nih.gov/articles/PMC12133834/"),
        ("S4", "International seroprevalence survey of IgE sensitisation to Dermatophagoides farinae in atopic dogs", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9292188/"),
        ("S5", "Allergens in veterinary medicine", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4716287/"),
        ("S6", "Common food allergen sources in dogs and cats", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4710035/"),
        ("S7", "Can in vivo or in vitro tests diagnose adverse food reactions in dogs and cats?", "https://link.springer.com/article/10.1186/s12917-017-1142-0"),
        ("S8", "Cutaneous manifestations of adverse food reactions in dogs and cats", "https://link.springer.com/article/10.1186/s12917-019-1880-2"),
        ("S9", "MSD Veterinary Manual - Feline atopic dermatitis", "https://www.msdvetmanual.com/integumentary-system/atopic-dermatitis/feline-atopic-dermatitis"),
        ("S10", "Review of feline allergic skin disease", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5606602/"),
    ]
    add_table(doc, ["ID", "Source", "URL"], sources, [650, 3600, 5110])

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    make_doc()
